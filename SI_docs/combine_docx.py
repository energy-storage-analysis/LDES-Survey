from docxcompose.composer import Composer
from docx import Document

master = Document("docx_written/TOC.docx")
composer = Composer(master)

# composer
SI_md = Document("output/SI_md.docx")
composer.append(SI_md)

# doc1 = Document("docx_written/LCOS.docx")
# composer.append(doc1)

# doc2 = Document("docx_written/EnergyForms.docx")
# composer.append(doc2)

composer.save("output/SI Automated.docx")